import os
import re
import logging
import tempfile
from typing import List, Dict, Any, Tuple

logger = logging.getLogger(__name__)
import pypdf
import docx
from PIL import Image

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import pytesseract
except Exception:
    pytesseract = None

from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from config import LLM_MODEL
from ollama_client import call_ollama


def clean_text_content(text: str) -> str:
    """Removes null bytes, non-printable control characters, and excess whitespace."""
    if not text:
        return ""
    cleaned = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", text)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    return cleaned.strip()


class FallbackEmbeddings:
    """100% Zero-Dependency In-Memory Deterministic Vector Embeddings (Never crashes)."""
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        return [self._hash_text(t) for t in texts]

    def embed_query(self, text: str) -> List[float]:
        return self._hash_text(text)

    def _hash_text(self, text: str, dim: int = 384) -> List[float]:
        import hashlib
        tokens = text.lower().split()
        vec = [0.0] * dim
        for i, t in enumerate(tokens):
            h = int(hashlib.md5(t.encode("utf-8")).hexdigest(), 16)
            vec[h % dim] += 1.0 / (i + 1)
        norm = sum(x * x for x in vec) ** 0.5
        return [x / norm if norm > 0 else 0.0 for x in vec]


class RAGEngine:
    def __init__(self, llm_model: str = "llama3.2:1b", embed_model: str = "all-MiniLM-L6-v2"):
        self.llm_model = llm_model
        
        # Robust Multi-Tier Embeddings (HuggingFace -> OllamaEmbeddings -> Zero-Dependency Fallback)
        try:
            from langchain_community.embeddings import HuggingFaceEmbeddings
            self.embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        except Exception:
            try:
                from langchain_ollama import OllamaEmbeddings
                self.embeddings = OllamaEmbeddings(model="nomic-embed-text")
            except Exception:
                self.embeddings = FallbackEmbeddings()

        self.vector_store = None
        self.all_documents: List[Document] = []
        self.doc_metadata: Dict[str, Any] = {}

    def extract_text_from_pdf(self, file_path: str, filename: str, max_pages_limit: int = 999999) -> List[Document]:
        """Multi-stage high-speed PDF text extraction indexing all document pages in seconds."""
        docs = []

        # Method 1: Fast pypdf extraction (Reads 100+ pages in < 1 second)
        try:
            reader = pypdf.PdfReader(file_path)
            num_pages = len(reader.pages)
            process_pages = min(num_pages, max_pages_limit)

            for idx in range(process_pages):
                page = reader.pages[idx]
                page_num = idx + 1
                text = clean_text_content(page.extract_text() or "")

                if not text:
                    text = f"[Page {page_num}: Visual / Diagram content]"

                docs.append(Document(
                    page_content=text,
                    metadata={
                        "source": filename,
                        "page": page_num,
                        "total_pages": num_pages
                    }
                ))

            if docs and sum(len(d.page_content.split()) for d in docs) > 30:
                return docs
        except Exception:
            docs = []

        # Method 2: pdfplumber fallback
        if pdfplumber is not None:
            try:
                with pdfplumber.open(file_path) as pdf:
                    num_pages = len(pdf.pages)
                    process_pages = min(num_pages, max_pages_limit)

                    for idx in range(process_pages):
                        page = pdf.pages[idx]
                        page_num = idx + 1
                        page_text = page.extract_text(layout=False) or page.extract_text() or ""
                        page_text = clean_text_content(page_text)

                        if not page_text:
                            page_text = f"[Page {page_num}: Visual / Diagram content]"

                        docs.append(Document(
                            page_content=page_text,
                            metadata={
                                "source": filename,
                                "page": page_num,
                                "total_pages": num_pages
                            }
                        ))
            except Exception:
                pass

        return docs

    def extract_text_from_docx(self, file_path: str, filename: str) -> List[Document]:
        """Extract text from DOCX file."""
        doc_obj = docx.Document(file_path)
        paragraphs = [clean_text_content(p.text) for p in doc_obj.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        
        page_size = 1800
        total_pages = max(1, (len(full_text) + page_size - 1) // page_size)
        docs = []
        
        for p_idx in range(total_pages):
            start = p_idx * page_size
            end = min(start + page_size, len(full_text))
            chunk_text = full_text[start:end]
            docs.append(Document(
                page_content=chunk_text,
                metadata={
                    "source": filename,
                    "page": p_idx + 1,
                    "total_pages": total_pages
                }
            ))
        return docs

    def extract_text_from_txt(self, file_path: str, filename: str) -> List[Document]:
        """Extract text from TXT file."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = clean_text_content(f.read())
            
        page_size = 2000
        total_pages = max(1, (len(text) + page_size - 1) // page_size)
        docs = []
        
        for p_idx in range(total_pages):
            start = p_idx * page_size
            end = min(start + page_size, len(text))
            chunk_text = text[start:end]
            docs.append(Document(
                page_content=chunk_text,
                metadata={
                    "source": filename,
                    "page": p_idx + 1,
                    "total_pages": total_pages
                }
            ))
        return docs

    def extract_text_from_image(self, file_path: str, filename: str) -> List[Document]:
        """Extract text from scanned image using pytesseract."""
        text = ""
        if pytesseract is not None:
            try:
                img = Image.open(file_path)
                text = clean_text_content(pytesseract.image_to_string(img))
            except Exception as e:
                text = f"[OCR Extraction Error: {str(e)}]"
        else:
            text = "[Tesseract OCR engine is not configured.]"
            
        return [Document(
            page_content=text if text else "[Image containing no readable text]",
            metadata={"source": filename, "page": 1, "total_pages": 1}
        )]

    def process_uploaded_file(self, file_bytes: bytes, filename: str, max_size_mb: float = 20.0, **kwargs) -> Tuple[List[Document], Dict[str, Any]]:
        """Processes an uploaded file, enforcing 20MB size limit for local Ollama llama3.2:1b processing."""
        max_limit = kwargs.get("max_size_mb", max_size_mb)
        file_size_mb = len(file_bytes) / (1024 * 1024)
        if file_size_mb > max_limit:
            raise ValueError(f"File size ({file_size_mb:.1f} MB) exceeds the {max_limit} MB limit for local Ollama processing.")

        ext = os.path.splitext(filename)[1].lower()
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            if ext == ".pdf":
                raw_docs = self.extract_text_from_pdf(tmp_path, filename, max_pages_limit=999999)
            elif ext in [".docx", ".doc"]:
                raw_docs = self.extract_text_from_docx(tmp_path, filename)
            elif ext in [".txt"]:
                raw_docs = self.extract_text_from_txt(tmp_path, filename)
            elif ext in [".png", ".jpg", ".jpeg"]:
                raw_docs = self.extract_text_from_image(tmp_path, filename)
            else:
                raise ValueError(f"Unsupported file format: {ext}")
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)

        total_pages = raw_docs[0].metadata.get("total_pages", len(raw_docs)) if raw_docs else 0
        total_words = sum(len(d.page_content.split()) for d in raw_docs)
        total_chars = sum(len(d.page_content) for d in raw_docs)

        meta = {
            "filename": filename,
            "total_pages": total_pages,
            "indexed_pages": len(raw_docs),
            "total_words": total_words,
            "total_chars": total_chars,
            "size_mb": round(file_size_mb, 2)
        }
        return raw_docs, meta

    def build_vector_store(self, raw_documents: List[Document]):
        """Splits raw documents into concise 400-char chunks using fast in-memory HuggingFace embeddings."""
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=400,
            chunk_overlap=60,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        chunks = splitter.split_documents(raw_documents)
        self.all_documents = raw_documents
        
        # Build FAISS store using local embeddings with automatic crash-proof fallback
        try:
            self.vector_store = FAISS.from_documents(chunks, self.embeddings)
        except Exception:
            logger.warning("Primary embedding provider encountered connection/runtime error. Falling back to crash-proof FallbackEmbeddings.")
            self.embeddings = FallbackEmbeddings()
            self.vector_store = FAISS.from_documents(chunks, self.embeddings)

        return len(chunks)

    def query(self, user_query: str, top_k: int = 3) -> Dict[str, Any]:
        """Queries RAG index with top 3 small chunks for high accuracy with Ollama 3.2:1b."""
        if not self.vector_store:
            return {"answer": "No documents indexed yet. Please upload a document first.", "sources": []}

        docs_and_scores = self.vector_store.similarity_search_with_score(user_query, k=top_k)
        
        context_parts = []
        sources = []
        
        for idx, (doc, score) in enumerate(docs_and_scores):
            src_file = doc.metadata.get("source", "Unknown Document")
            page_num = doc.metadata.get("page", "?")
            snippet = doc.page_content.strip()
            
            context_parts.append(f"--- Snippet [{idx+1}] (Page {page_num}) ---\n{snippet}")
            sources.append({
                "id": idx + 1,
                "source": src_file,
                "page": page_num,
                "snippet": snippet,
                "score": float(score)
            })

        context_str = "\n\n".join(context_parts)
        
        prompt = f"""You are EDUMIND Document Brain. Answer the question accurately using ONLY the context snippets below.
Always include exact citations in format `[Source: <filename>, Page: <page_number>]`.

Context Excerpts:
{context_str}

Question: {user_query}

Answer (concise & cited):"""

        answer = call_ollama(prompt, model=self.llm_model)
        return {
            "answer": answer,
            "sources": sources
        }
